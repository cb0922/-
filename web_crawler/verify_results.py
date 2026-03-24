#!/usr/bin/env python3
"""
验证测试结果
"""
import os
from docx import Document
import json

print("=" * 60)
print("测试结果验证")
print("=" * 60)

# 1. 检查Word文档
word_dir = 'test_optimized_output/word_reports'
if os.path.exists(word_dir):
    files = [f for f in os.listdir(word_dir) if f.endswith('.docx')]
    if files:
        word_path = os.path.join(word_dir, files[0])
        doc = Document(word_path)
        print(f"\n1. Word文档: {files[0]}")
        print(f"   段落数: {len(doc.paragraphs)}")
        
        # 统计长段落
        long_paras = [p for p in doc.paragraphs if len(p.text) > 200]
        print(f"   正文段落(>200字): {len(long_paras)}")
        if long_paras:
            print(f"   最长段落: {max(len(p.text) for p in long_paras)} 字符")
    else:
        print("\n1. Word文档: 未找到")
else:
    print("\n1. Word文档目录: 不存在")

# 2. 检查JSON数据
json_dir = 'test_optimized_output/data'
if os.path.exists(json_dir):
    files = [f for f in os.listdir(json_dir) if f.endswith('.json')]
    if files:
        json_path = os.path.join(json_dir, files[0])
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        print(f"\n2. JSON数据: {files[0]}")
        print(f"   记录数: {len(data)}")
        if data:
            # 检查text字段
            texts = [d.get('text', '') for d in data if d.get('text')]
            print(f"   有text字段: {len(texts)}/{len(data)}")
            if texts:
                avg_len = sum(len(t) for t in texts) / len(texts)
                print(f"   平均内容长度: {avg_len:.0f} 字符")
    else:
        print("\n2. JSON数据: 未找到")
else:
    print("\n2. JSON数据目录: 不存在")

# 3. 检查PDF
pdf_dirs = ['pdfs/test_large', 'pdfs/competitions', 'pdfs/enhanced']
pdf_count = 0
for d in pdf_dirs:
    if os.path.exists(d):
        files = [f for f in os.listdir(d) if f.endswith('.pdf')]
        pdf_count += len(files)

print(f"\n3. PDF文件: {pdf_count} 个")

# 4. 总结
print("\n" + "=" * 60)
print("验证结果:")
print("=" * 60)
if long_paras and len(data) > 0 and pdf_count > 0:
    print("所有功能正常!")
    print(f"Word文档包含 {len(long_paras)} 个正文段落")
    print(f"JSON包含 {len(data)} 条记录")
    print(f"下载了 {pdf_count} 个PDF文件")
else:
    print("部分功能可能有问题")
print("=" * 60)
