#!/usr/bin/env python3
"""
Word文档生成器
将抓取的通知/公告内容整理成Word文档
"""
import os
import re
from datetime import datetime
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


class NoticeWordGenerator:
    """通知/公告 Word 文档生成器"""
    
    # 颜色定义
    COLOR_BLUE = RGBColor(0x25, 0x63, 0xeb)  # 主蓝色
    COLOR_DARK_BLUE_GRAY = RGBColor(0x2c, 0x3e, 0x50)  # 深蓝灰色
    COLOR_GRAY = RGBColor(0x5d, 0x6d, 0x7e)  # 灰色
    COLOR_LIGHT_GRAY = RGBColor(0x64, 0x70, 0x8b)  # 浅灰色
    COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)  # 黑色
    COLOR_RED = RGBColor(0xef, 0x44, 0x44)  # 红色
    COLOR_TABLE_HEADER = RGBColor(0xd6, 0xea, 0xf8)  # 浅蓝色表头
    
    def __init__(self, output_dir: str = "./word_reports"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def generate(self, data: List[Dict], title: str = "通知公告汇总") -> str:
        """
        生成Word文档
        :param data: 爬取的数据列表
        :param title: 文档标题
        :return: 生成的文件路径
        """
        doc = Document()
        
        # 设置页面样式（页边距）
        self._set_page_margins(doc)
        
        # 设置文档默认字体和样式
        self._set_document_style(doc)
        
        # 设置标题样式
        self._set_heading_styles(doc)
        
        # 添加页眉页脚
        self._add_header_footer(doc, title)
        
        # 添加封面/标题
        self._add_title(doc, title)
        
        # 添加统计信息
        self._add_statistics(doc, data)
        
        # 添加目录
        self._add_toc(doc, data)
        
        # 添加详细内容
        self._add_content_details(doc, data)
        
        # 保存文档
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"通知公告汇总_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        doc.save(filepath)
        print(f"  [WORD] Word文档已保存: {filepath}")
        return filepath
    
    def _set_page_margins(self, doc: Document):
        """设置页边距（上下2.54cm，左右3.17cm）"""
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)
    
    def _set_document_style(self, doc: Document):
        """设置文档默认样式"""
        # 设置Normal样式
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 设置段落格式
        style.paragraph_format.line_spacing = 1.5  # 1.5倍行距
        style.paragraph_format.first_line_indent = Inches(0.4)  # 首行缩进2字符（约0.4英寸）
        style.paragraph_format.space_before = Pt(0)  # 段前0
        style.paragraph_format.space_after = Pt(6)  # 段后6pt
    
    def _set_heading_styles(self, doc: Document):
        """设置标题样式"""
        # 一级标题：16号字，加粗，黑色
        heading1 = doc.styles['Heading 1']
        heading1.font.name = 'Microsoft YaHei'
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        heading1.font.color.rgb = self.COLOR_BLACK
        heading1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        heading1.paragraph_format.space_before = Pt(12)
        heading1.paragraph_format.space_after = Pt(6)
        heading1.paragraph_format.first_line_indent = Inches(0)
        
        # 二级标题：14号字，加粗，深蓝灰色
        heading2 = doc.styles['Heading 2']
        heading2.font.name = 'Microsoft YaHei'
        heading2.font.size = Pt(14)
        heading2.font.bold = True
        heading2.font.color.rgb = self.COLOR_DARK_BLUE_GRAY
        heading2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        heading2.paragraph_format.space_before = Pt(10)
        heading2.paragraph_format.space_after = Pt(6)
        heading2.paragraph_format.first_line_indent = Inches(0)
        
        # 三级标题：12号字，加粗，灰色
        heading3 = doc.styles['Heading 3']
        heading3.font.name = 'Microsoft YaHei'
        heading3.font.size = Pt(12)
        heading3.font.bold = True
        heading3.font.color.rgb = self.COLOR_GRAY
        heading3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        heading3.paragraph_format.space_before = Pt(8)
        heading3.paragraph_format.space_after = Pt(6)
        heading3.paragraph_format.first_line_indent = Inches(0)
    
    def _add_header_footer(self, doc: Document, title: str):
        """添加页眉和页脚"""
        section = doc.sections[0]
        
        # 添加页眉
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = title
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header_para.runs:
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(9)
            run.font.color.rgb = self.COLOR_LIGHT_GRAY
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 添加页脚（页码）
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加页码域
        self._add_page_number(footer_para)
    
    def _add_page_number(self, paragraph):
        """在段落中添加页码"""
        run = paragraph.add_run()
        
        # 创建fldChar元素
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
        
        # 添加" / "
        run2 = paragraph.add_run(" / ")
        run2.font.name = 'Microsoft YaHei'
        run2.font.size = Pt(9)
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 添加总页数
        run3 = paragraph.add_run()
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'begin')
        
        instrText2 = OxmlElement('w:instrText')
        instrText2.set(qn('xml:space'), 'preserve')
        instrText2.text = "NUMPAGES"
        
        fldChar5 = OxmlElement('w:fldChar')
        fldChar5.set(qn('w:fldCharType'), 'separate')
        
        fldChar6 = OxmlElement('w:fldChar')
        fldChar6.set(qn('w:fldCharType'), 'end')
        
        run3._r.append(fldChar4)
        run3._r.append(instrText2)
        run3._r.append(fldChar5)
        run3._r.append(fldChar6)
        
        for r in [run, run3]:
            r.font.name = 'Microsoft YaHei'
            r.font.size = Pt(9)
            r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    def _add_title(self, doc: Document, title: str):
        """添加文档标题"""
        # 主标题：22号字，加粗，居中，蓝色
        heading = doc.add_heading(level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = heading.add_run(title)
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = self.COLOR_BLUE
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # 添加生成时间
        time_para = doc.add_paragraph()
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        time_para.paragraph_format.first_line_indent = Inches(0)  # 时间居中，无缩进
        time_run = time_para.add_run(f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}")
        time_run.font.name = 'Microsoft YaHei'
        time_run.font.size = Pt(10)
        time_run.font.color.rgb = self.COLOR_LIGHT_GRAY
        time_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        doc.add_paragraph()  # 空行
    
    def _add_statistics(self, doc: Document, data: List[Dict]):
        """添加统计信息"""
        doc.add_heading("一、数据统计", level=1)
        
        # 统计各类型数量
        total = len(data)
        with_pdf = sum(1 for item in data if item.get('competition_pdfs') or item.get('competition_documents'))
        with_notice = sum(1 for item in data if self._is_notice_content(item))
        
        stats = [
            f"总共抓取网页：{total} 个",
            f"包含文档附件：{with_pdf} 个",
            f"包含通知公告：{with_notice} 个",
        ]
        
        for stat in stats:
            p = doc.add_paragraph(stat, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)  # 列表项缩进0.5英寸
            p.paragraph_format.first_line_indent = Inches(-0.2)  # 悬挂缩进
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(11)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        doc.add_paragraph()  # 空行
    
    def _add_toc(self, doc: Document, data: List[Dict]):
        """添加目录"""
        doc.add_heading("二、内容目录", level=1)
        
        for i, item in enumerate(data, 1):
            title = item.get('title', '无标题')
            
            # 创建目录条目
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(4)
            
            # 序号和标题
            run = p.add_run(f"{i}. {title}")
            run.font.size = Pt(11)
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            
            # 如果有PDF，标记出来
            if item.get('competition_pdfs') or item.get('competition_documents'):
                pdf_run = p.add_run(" [有附件]")
                pdf_run.font.color.rgb = self.COLOR_RED
                pdf_run.font.size = Pt(9)
                pdf_run.font.name = 'Microsoft YaHei'
                pdf_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        doc.add_page_break()  # 分页
    
    def _add_content_details(self, doc: Document, data: List[Dict]):
        """添加详细内容（完整正文）"""
        doc.add_heading("三、详细内容", level=1)
        
        for i, item in enumerate(data, 1):
            # 提取通知/公告信息（完整内容）
            notice_info = self._extract_notice_info(item, full_content=True)
            
            # 添加小标题
            doc.add_heading(f"{i}. {notice_info['title']}", level=2)
            
            # 添加元信息表格
            self._add_metadata_table(doc, notice_info)
            
            # 添加完整正文内容
            content_structure = notice_info.get('content_structure', [])
            if content_structure:
                doc.add_heading("正文内容：", level=3)
                # 使用解析后的结构化内容
                self._add_structured_content(doc, content_structure)
            elif notice_info.get('full_content', ''):
                # 兼容旧格式，直接显示文本
                doc.add_heading("正文内容：", level=3)
                full_content = notice_info['full_content']
                paragraphs = full_content.split('\n')
                for para_text in paragraphs:
                    para_text = para_text.strip()
                    if para_text:
                        content_para = doc.add_paragraph(para_text)
                        content_para.paragraph_format.first_line_indent = Inches(0.4)
                        content_para.paragraph_format.space_after = Pt(6)
                        content_para.paragraph_format.line_spacing = 1.5
            
            # 添加PDF附件列表
            if notice_info['pdfs']:
                doc.add_heading("附件列表：", level=3)
                for pdf in notice_info['pdfs']:
                    pdf_para = doc.add_paragraph(style='List Bullet')
                    pdf_para.paragraph_format.left_indent = Inches(0.5)
                    pdf_para.paragraph_format.first_line_indent = Inches(-0.2)
                    
                    pdf_name = pdf.get('filename', '未命名.pdf')
                    pdf_run = pdf_para.add_run(f"📄 {pdf_name}")
                    pdf_run.font.size = Pt(10)
                    pdf_run.font.name = 'Microsoft YaHei'
                    pdf_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            
            # 添加分隔线（使用美观的分隔符）
            if i < len(data):
                separator_para = doc.add_paragraph()
                separator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                separator_para.paragraph_format.first_line_indent = Inches(0)
                separator_run = separator_para.add_run("━━━" * 15)
                separator_run.font.color.rgb = self.COLOR_LIGHT_GRAY
                separator_run.font.size = Pt(8)
                
                # 分页
                doc.add_page_break()
    
    def _add_structured_content(self, doc: Document, content_structure: List[Dict]):
        """添加结构化的正文内容"""
        for item in content_structure:
            item_type = item.get('type', 'paragraph')
            content = item.get('content', '').strip()
            
            if not content:
                continue
            
            if item_type == 'heading':
                # 标题
                level = item.get('level', 1)
                if level == 1:
                    p = doc.add_heading(content, level=4)
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(content)
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = self.COLOR_DARK_BLUE_GRAY
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                    p.paragraph_format.first_line_indent = Inches(0)
                    p.paragraph_format.space_before = Pt(8)
                    p.paragraph_format.space_after = Pt(4)
            
            elif item_type == 'list_number':
                # 数字列表
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                run = p.add_run(content)
                run.font.size = Pt(11)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            
            elif item_type == 'list_bullet':
                # 项目符号列表
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.2)
                # 移除原文中的项目符号，因为样式会自动添加
                clean_content = re.sub(r'^[•·\-\*]\s*', '', content)
                run = p.add_run(clean_content)
                run.font.size = Pt(11)
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            
            elif item_type == 'paragraph':
                # 普通段落
                p = doc.add_paragraph(content)
                p.paragraph_format.first_line_indent = Inches(0.4)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.5
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    def _add_metadata_table(self, doc: Document, info: Dict):
        """添加元信息表格"""
        # 创建表格
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置列宽（第一列30%，第二列70%）
        table.allow_autofit = False
        table.columns[0].width = Inches(1.5)
        table.columns[1].width = Inches(3.5)
        
        # 填充数据
        metadata = [
            ("来源网址", info['url']),
            ("发布时间", info.get('publish_time', '未知')),
            ("关键词", ", ".join(info.get('keywords', [])) or "无"),
        ]
        
        for i, (key, value) in enumerate(metadata):
            row = table.rows[i]
            
            # 设置单元格背景色（浅蓝色表头）
            cell_key = row.cells[0]
            cell_key.text = key
            
            # 设置表头单元格背景色
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'D6EAF8')
            cell_key._tc.get_or_add_tcPr().append(shading_elm)
            
            # 设置值单元格
            cell_value = row.cells[1]
            display_value = value or ""
            if len(display_value) > 100:
                display_value = display_value[:100] + "..."
            cell_value.text = display_value
            
            # 设置单元格样式
            for j, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.first_line_indent = Inches(0)
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Microsoft YaHei'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                        # 表头加粗
                        if j == 0:
                            run.font.bold = True
        
        doc.add_paragraph()  # 空行
    
    def _extract_notice_info(self, item: Dict, full_content: bool = False) -> Dict:
        """提取通知/公告信息
        :param full_content: 是否提取完整内容（Word文档用）
        """
        # 优先使用已提取的text，如果没有则尝试从content解析
        text_content = item.get('text', '')
        html_content = item.get('content', '')
        
        # 如果text为空，尝试从HTML提取
        if not text_content and html_content and html_content != '[HTML_CONTENT_TRUNCATED]':
            try:
                from core.parser import ContentParser
                parser = ContentParser(html_content)
                text_content = parser.get_main_content()
            except Exception:
                pass
        
        # 使用text_content作为主要内容
        content = text_content or html_content or ''
        
        # 解析HTML结构（如果有）
        content_structure = []
        if html_content and html_content != '[HTML_CONTENT_TRUNCATED]':
            content_structure = self._parse_html_structure(html_content)
        
        # 如果结构解析失败，回退到文本分段
        if not content_structure and content:
            content_structure = self._parse_text_structure(content)
        
        # 尝试提取时间
        publish_time = self._extract_date(content)
        
        # 提取关键词
        keywords = self._extract_keywords(content)
        
        result = {
            'title': item.get('title', '无标题'),
            'url': item.get('url', ''),
            'content': content[:500] + "..." if len(content) > 500 else content,
            'full_content': content if full_content else '',
            'content_structure': content_structure if full_content else [],
            'publish_time': publish_time,
            'keywords': keywords,
            'pdfs': item.get('competition_documents') or item.get('competition_pdfs', [])
        }
        
        return result
    
    def _parse_html_structure(self, html_content: str) -> List[Dict]:
        """从HTML中解析内容结构，保留段落、标题、列表等"""
        structure = []
        
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 移除script和style标签
            for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
                tag.decompose()
            
            # 找到主要内容区域
            main_content = None
            for selector in ['article', 'main', '.content', '#content', '.article', '.post', '.entry']:
                main_content = soup.select_one(selector)
                if main_content:
                    break
            
            if not main_content:
                main_content = soup.body or soup
            
            # 遍历所有相关元素
            for elem in main_content.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'div']):
                text = elem.get_text(strip=True)
                if not text or len(text) < 3:
                    continue
                
                tag_name = elem.name.lower()
                
                # 识别标题
                if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    level = int(tag_name[1])
                    structure.append({
                        'type': 'heading',
                        'level': level,
                        'content': text
                    })
                
                # 识别列表项
                elif tag_name == 'li':
                    # 检查父元素是ul还是ol
                    parent = elem.find_parent(['ul', 'ol'])
                    if parent and parent.name == 'ol':
                        structure.append({
                            'type': 'list_number',
                            'content': text
                        })
                    else:
                        structure.append({
                            'type': 'list_bullet',
                            'content': text
                        })
                
                # 识别段落
                elif tag_name == 'p':
                    structure.append({
                        'type': 'paragraph',
                        'content': text
                    })
                
                # 处理div（可能包含其他内容）
                elif tag_name == 'div':
                    # 检查是否是纯文本div
                    children = list(elem.children)
                    if all(isinstance(c, str) or c.name in ['br', 'span'] for c in children):
                        structure.append({
                            'type': 'paragraph',
                            'content': text
                        })
            
        except ImportError:
            # 如果没有BeautifulSoup，回退到文本解析
            pass
        except Exception as e:
            print(f"  [WARN] HTML解析失败: {e}")
        
        return structure
    
    def _parse_text_structure(self, text_content: str) -> List[Dict]:
        """从纯文本中解析内容结构"""
        structure = []
        lines = text_content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 检测标题（以"："或":"结尾的短文本）
            if len(line) < 50 and (line.endswith('：') or line.endswith(':')):
                structure.append({
                    'type': 'heading',
                    'level': 2,
                    'content': line
                })
            
            # 检测数字列表（以数字和点开头）
            elif re.match(r'^\d+[\.、]\s*', line):
                structure.append({
                    'type': 'list_number',
                    'content': line
                })
            
            # 检测项目符号列表
            elif re.match(r'^[•·\-\*]\s*', line):
                structure.append({
                    'type': 'list_bullet',
                    'content': line
                })
            
            # 普通段落
            else:
                structure.append({
                    'type': 'paragraph',
                    'content': line
                })
        
        return structure
    
    def _extract_date(self, text: str) -> Optional[str]:
        """从文本中提取日期"""
        if not text:
            return None
        
        # 匹配常见日期格式
        patterns = [
            r'(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日',
            r'(\d{4})-(\d{2})-(\d{2})',
            r'(\d{4})/(\d{2})/(\d{2})',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                groups = match.groups()
                if len(groups) == 3:
                    return f"{groups[0]}年{groups[1]}月{groups[2]}日"
        
        return None
    
    def _extract_keywords(self, text: str) -> List[str]:
        """提取关键词"""
        if not text:
            return []
        
        # 关键词列表
        keywords = [
            "大赛", "比赛", "竞赛", "评职", "评奖", "评优", "评先",
            "微课", "公开课", "优质课", "精品课", "征文", "论文",
            "通知", "公告", "公示", "报名", "参赛"
        ]
        
        found = []
        text_lower = text.lower()
        for kw in keywords:
            if kw in text_lower and kw not in found:
                found.append(kw)
        
        return found[:5]  # 最多返回5个
    
    def _is_notice_content(self, item: Dict) -> bool:
        """判断是否包含通知/公告内容"""
        title = item.get('title', '').lower()
        text = (item.get('text', '') or item.get('content', '')).lower()
        
        notice_keywords = ['通知', '公告', '公示', '简章', '方案']
        return any(kw in title or kw in text for kw in notice_keywords)


def generate_word_report(crawl_results: List[Dict], output_dir: str = "./word_reports") -> str:
    """
    便捷的Word报告生成函数
    """
    generator = NoticeWordGenerator(output_dir)
    return generator.generate(crawl_results, title="通知公告汇总报告")


if __name__ == "__main__":
    # 测试
    test_data = [
        {
            "title": "第九届全国语文微课大赛通知",
            "url": "http://example.com/1",
            "text": "关于举办第九届全国语文微课大赛的通知\n\n一、大赛宗旨\n为提高语文教师的教学能力，推动语文教育创新发展，特举办本次大赛。\n\n二、参赛对象\n1. 全国中小学语文教师\n2. 高校语文教育专业师生\n3. 语文教研员\n\n三、比赛要求\n• 微课时长：5-10分钟\n• 内容要求：体现新课标理念\n• 技术要求：画面清晰，声音清楚",
            "competition_pdfs": [
                {"filename": "大赛通知.pdf"},
                {"filename": "报名表.docx"}
            ]
        },
        {
            "title": "2024年教师职称评审公告",
            "url": "http://example.com/2",
            "text": "2024年教师职称评审公告\n\n根据教育部文件精神，现将2024年职称评审有关事项公告如下：\n\n一、评审条件\n1. 具有相应教师资格证书\n2. 完成规定学时培训\n3. 近五年年度考核合格以上\n\n二、申报材料\n• 个人述职报告\n• 教学成果证明\n• 获奖证书复印件",
            "competition_pdfs": []
        }
    ]
    
    filepath = generate_word_report(test_data)
    print(f"测试文档已生成: {filepath}")
